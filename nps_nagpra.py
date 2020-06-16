import requests
from bs4 import BeautifulSoup
import os
from docx.api import Document
import glob
import subprocess
import json
import csv
import codecs
from collections import defaultdict


def get_source(uri):
    r = requests.get(uri)
    if r.status_code == requests.codes.ok:
        source = r.content
        return source
    else:
        return


def get_links(doc, base):
    soup = BeautifulSoup(doc, features="lxml")
    links = ["".join([base, link['href']]) for link in soup.find_all('a', href=True)]
    return links


def filter_links(links, filter_str):
    return [link for link in links if filter_str in link]


def fetch(url_list, destination):
    for url in url_list:
        filename = url.split("/")[-1]
        destination_file = os.path.join(destination, filename)
        print(destination_file)
        r = requests.get(url, allow_redirects=True)
        with open(destination_file, "wb") as f:
            f.write(r.content)


def convert_to_docx(word_doc):
    # doc = Document(word_doc)
    subprocess.call(['/Applications/LibreOffice.app/Contents/MacOS/soffice',
                     '--headless', '--convert-to', 'docx', word_doc])


def convert_to_csv(word_doc):
    # doc = Document(word_doc)
    subprocess.call(['python',
                     'bin/docx2csv', '--format', 'csv', word_doc])


def normalize_csv(csv_file):
    print(csv_file)
    with codecs.open(csv_file, "r", encoding="utf-8") as f:
        c = csv.DictReader(f)
        rows = []
        for row in c:
            print(row)
            row_dict = {k[2:-1]:v[2:-1] for k, v in row.items()}
            rows.append(row_dict)
        # print(rows[0])
        fieldnames = [k for k, _ in rows[0].items()]
        headers = {n:n for n in fieldnames}
        with open(csv_file.replace("_1.csv", "_2.csv"), "w") as f_out:
            dw = csv.DictWriter(f_out, fieldnames=fieldnames)
            dw.writerow(headers)
            for crow in rows:
                dw.writerow(crow)
        # print(json.dumps(rows, indent=4))


def iter_visual_cells(row):
    """
    https://github.com/python-openxml/python-docx/issues/344#issuecomment-271390490

    :param row:
    :return:
    """
    prior_tc = None
    merged_row = []
    for cell in row.cells:
        this_tc = cell._tc
        if this_tc is not prior_tc:  # skip cells pointing to same `<w:tc>` element
            merged_row.append(cell.text)
        prior_tc = this_tc
    return merged_row



def tabulate(word_doc):
    print(word_doc)
    json_file = word_doc.replace(".docx", ".json")
    if not os.path.exists(json_file):
        print(f"Loading {word_doc} using python-docx.")
        doc = Document(word_doc)
        print("Extracting tables using python-docx.")
        table = doc.tables[0]
        keys = ['MAP #', 'MAP Name', 'State', 'County', 'Tribe Named in Treaty', 'Present-Day Tribe']
        big_list = []
        print("Iterating rows to handle merged cells.")
        len_rows = len(table.rows)
        for n, row in enumerate(table.rows):
            # print(f"{n} row of {len_rows}")
            foo = iter_visual_cells(row)
            big_list.extend(foo)
        final = []
        total = len(big_list)
        print(f"Beginning to iterate and chunk the list")
        for i, x in enumerate(big_list):
            print(f"{i} of {total}")
            if all([s.isdigit() for s in x.strip()]) and x is not '':
                final.append(dict(zip(keys, big_list[i:i + 6])))
        with open(json_file, "w") as fout:
            print(f"Writing {json_file}")
            json.dump(final, fout, indent=4)


def tribes():
    docs = glob.glob("./nps_docs/*.json")
    historic = defaultdict(list)
    present = defaultdict(list)
    for doc in docs:
        print(doc)
        with open(doc, "r") as f:
            j = json.load(f)
            for d in j:
                # Fix the messed up encoding from the Microsoft docx conversion
                key = bytes([ord(char) for char in d["Tribe Named in Treaty"]]).decode('latin-1')
                val = bytes([ord(char) for char in d["Present-Day Tribe"]]).decode('latin-1')
                historic[key].append(val)
                present[val].append(key)
    for k, v in historic.items():
        historic[k] = list(set(v))
    for k, v in present.items():
        present[k] = list(set(v))
    with open("historic.json", "w") as f_out:
        json.dump(historic, f_out, indent=2, sort_keys=True)
    with open("presentday.json", "w") as p_out:
        json.dump(present, p_out, indent=2, sort_keys=True)


if __name__ == "__main__":
    # docs = glob.glob("/Volumes/Seagate Expansion "
    #                  "Drive/Github/ida-treaty-explorer-data/nps_docs/*.docx")
    # for doc in docs:
    #     tabulate(doc)
    tribes()



